import PyPDF2
import pdfplumber
import pandas as pd
import docx
from langdetect import detect
from app.services.ocr_service import OCRService
from app.models.schemas import FileType, DocumentMetadata
from typing import Dict, Any
import os

class FileProcessor:
    def __init__(self):
        self.ocr_service = OCRService()
    
    def process_file(self, file_path: str, file_type: FileType, 
                     doc_id: str, filename: str) -> Dict[str, Any]:
        """Main processing dispatcher"""
        
        if file_type == FileType.PDF:
            return self._process_pdf(file_path, doc_id, filename)
        elif file_type == FileType.IMAGE:
            return self._process_image(file_path, doc_id, filename)
        elif file_type == FileType.TEXT:
            return self._process_text(file_path, doc_id, filename)
        elif file_type == FileType.CSV:
            return self._process_csv(file_path, doc_id, filename)
        elif file_type == FileType.EXCEL:
            return self._process_excel(file_path, doc_id, filename)
        elif file_type == FileType.DOCX:
            return self._process_docx(file_path, doc_id, filename)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _process_pdf(self, file_path: str, doc_id: str, 
                     filename: str) -> Dict[str, Any]:
        """Process PDF - try direct extraction, fall back to OCR"""
        
        # Try direct extraction first
        text = ""
        tables = []
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
                
                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    for table in page_tables:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        tables.append({
                            'page': page.page_number,
                            'data': df.to_dict('records')
                        })
        
        # If little text extracted, it's likely scanned
        if len(text.strip()) < 100:
            text = self.ocr_service.extract_from_pdf(file_path)
            page_count = text.count("--- Page")
        
        # Detect language
        language = self._detect_language(text)
        
        return {
            'content': text,
            'tables': tables,
            'metadata': {
                'document_id': doc_id,
                'filename': filename,
                'file_type': FileType.PDF,
                'language': language,
                'page_count': page_count
            }
        }
    
    def _process_image(self, file_path: str, doc_id: str, 
                      filename: str) -> Dict[str, Any]:
        """Process image with OCR"""
        text = self.ocr_service.extract_from_image(file_path)
        language = self._detect_language(text)
        
        return {
            'content': text,
            'tables': [],
            'metadata': {
                'document_id': doc_id,
                'filename': filename,
                'file_type': FileType.IMAGE,
                'language': language
            }
        }
    
    def _process_text(self, file_path: str, doc_id: str, 
                     filename: str) -> Dict[str, Any]:
        """Process plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        language = self._detect_language(text)
        
        return {
            'content': text,
            'tables': [],
            'metadata': {
                'document_id': doc_id,
                'filename': filename,
                'file_type': FileType.TEXT,
                'language': language
            }
        }
    
    def _process_csv(self, file_path: str, doc_id: str, 
                    filename: str) -> Dict[str, Any]:
        """Process CSV file"""
        df = pd.read_csv(file_path)
        
        # Convert to text representation
        text = f"CSV File: {filename}\n\n"
        text += f"Columns: {', '.join(df.columns.tolist())}\n\n"
        text += df.to_string(index=False)
        
        return {
            'content': text,
            'tables': [{'data': df.to_dict('records')}],
            'metadata': {
                'document_id': doc_id,
                'filename': filename,
                'file_type': FileType.CSV,
                'language': 'en',
                'row_count': len(df),
                'columns': df.columns.tolist()
            }
        }
    
    def _process_excel(self, file_path: str, doc_id: str, 
                      filename: str) -> Dict[str, Any]:
        """Process Excel file"""
        excel_file = pd.ExcelFile(file_path)
        
        text = f"Excel File: {filename}\n\n"
        tables = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            text += f"Sheet: {sheet_name}\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n"
            text += df.to_string(index=False) + "\n\n"
            
            tables.append({
                'sheet': sheet_name,
                'data': df.to_dict('records')
            })
        
        return {
            'content': text,
            'tables': tables,
            'metadata': {
                'document_id': doc_id,
                'filename': filename,
                'file_type': FileType.EXCEL,
                'language': 'en',
                'sheet_count': len(excel_file.sheet_names)
            }
        }
    
    def _process_docx(self, file_path: str, doc_id: str, 
                     filename: str) -> Dict[str, Any]:
        """Process Word document"""
        doc = docx.Document(file_path)
        
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract tables
        tables = []
        for table in doc.tables:
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            
            if data:
                df = pd.DataFrame(data[1:], columns=data[0])
                tables.append({'data': df.to_dict('records')})
        
        language = self._detect_language(text)
        
        return {
            'content': text,
            'tables': tables,
            'metadata': {
                'document_id': doc_id,
                'filename': filename,
                'file_type': FileType.DOCX,
                'language': language
            }
        }
    
    def _detect_language(self, text: str) -> str:
        """Detect text language"""
        try:
            if len(text.strip()) < 20:
                return 'unknown'
            return detect(text[:1000])  # Sample first 1000 chars
        except:
            return 'unknown'